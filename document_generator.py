import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from app import app

def generate_der_document(client):
    """Génère un Document d'Entrée en Relation (DER) pour le client en utilisant un modèle"""
    try:
        # Chemin vers le modèle DER
        template_path = os.path.join(os.path.dirname(__file__), 'templates_docs', 'der_template.docx')
        
        # Vérifier si le modèle existe
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Modèle DER non trouvé : {template_path}")
        
        # Charger le modèle
        doc = Document(template_path)
        
        # Préparer les données de remplacement
        replacements = {
            '{{date_entree_relation}}': client.date_entree_relation.strftime('%d/%m/%Y') if client.date_entree_relation else 'Non renseignée',
            '{{ville_client}}': client.ville if hasattr(client, 'ville') and client.ville else 'Non renseignée',
            '{{nom_client}}': client.nom or 'Non renseigné',
            '{{prenom_client}}': client.prenom or 'Non renseigné',
            '{{email_client}}': client.email or 'Non renseigné',
            '{{telephone_client}}': client.telephone or 'Non renseigné',
            '{{date_naissance_client}}': client.date_naissance.strftime('%d/%m/%Y') if client.date_naissance else 'Non renseignée',
            '{{adresse_client}}': client.adresse or 'Non renseignée',
            '{{profession_client}}': client.profession or 'Non renseignée'
        }
        
        # Remplacer les tags dans tous les paragraphes
        for paragraph in doc.paragraphs:
            for tag, value in replacements.items():
                if tag in paragraph.text:
                    paragraph.text = paragraph.text.replace(tag, value)
        
        # Remplacer les tags dans les tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for tag, value in replacements.items():
                        if tag in cell.text:
                            cell.text = cell.text.replace(tag, value)
        
        # Créer le dossier de destination s'il n'existe pas
        output_dir = os.path.join(app.root_path, 'generated_docs')
        os.makedirs(output_dir, exist_ok=True)
        
        # Nom du fichier de sortie
        filename = f"DER_{client.nom}_{client.prenom}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = os.path.join(output_dir, filename)
        
        # Sauvegarder le document personnalisé
        doc.save(output_path)
        
        return output_path
        
    except Exception as e:
        print(f"Erreur lors de la génération du DER : {str(e)}")
        return None
